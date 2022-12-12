from docx import Document
from docx.shared import Inches
import io
from django.http import HttpResponse

import pandas as pd

from .tables import add_row_table_reports, add_table_reports, add_spec_row_table_reports
from .models import Cabinets, PhDLDconnections, LogicDevices, DataObjects

def word_report(request, cab):

# СОЗДАЕМ ОТЧЕТ


    document = Document('base_types/templates/template.docx')
    document.add_heading(cab, 2)

    #cabinets = Cabinet.objects.all()
    #cabinets = cabinets.order_by('name')
    #for item in cabinets.iterator():
     #   add_row_table_reports(t,(item.name,'','','','','','',''))

    df = pd.DataFrame(columns=['signal', 'iec', 'attr_name', 'attr_status', 'alm', 'cus', 'rdu', 'ras', 'dataset'])
    datasets = set()

    # ищем шкаф в базе
    cabinet = Cabinets.objects.get(name=cab)


    if cabinet.terminal1: # если есть терминал 1 то добавляем таблицу
        document.add_heading('ИЭУ1: '+str(cabinet.terminal1), 4)


        #ищем состав лог. устройств в нем
        connections = PhDLDconnections.objects.all().filter(ied=cabinet.terminal1)

        for item in connections:
            print('---->', item.ld)
            ld_short_eng_name_field = str(item.ld)  # взяли имя для второго столбца ATCC/....
            ld = LogicDevices.objects.get(name=item.LDname)
            ld_short_rus_name_field = str(ld.short_name) # взяли имя для столбца сигнал в самой первой части ДЗАТ / ....
            print('ld_short_rus_name', ld_short_rus_name_field)

            lns = LDcontainer.objects.all().filter(name=item.LDname)
            for ln in lns:
                ln_short_rus_name_field = str(ln.short_fb_name) # взяли имя для столбца сигнал в самой первой части АРН Т / АВРС
                class_full_eng_name_field = str(ln.ln_prefix)+str(ln.ln_name)+ln.get_str_ln_instance

                print(ln.ln_type)
                das = LnTypes.objects.all().filter(name=ln.ln_type)
                for da in das:
                    #print('======================', da.data_obj)
                    data_obj_sen = DataObjects.objects.get(name=da.data_obj)
                    attr_rus_field = str(data_obj_sen.desc)
                    if da.signal:
                         attr_rus_field = str(da.signal)
                    _func_group = da.func_group # оставляем интегер для последующей фильтрации датасета по функ группе
                    _clue_attr_name_field = str(da.get_clue_attr) # готовые поля для Значащих атрибутов
                    _clue_attr_status_field = str(da.get_status) # готовые поля для Значащих атрибутов
                    _signal_field = ld_short_rus_name_field + ' / ' + ln_short_rus_name_field + ':' + attr_rus_field
                    _iec_name = ld_short_eng_name_field + '/' + class_full_eng_name_field + '.' + da.get_str_data_obj
                    _cus = da.get_cus
                    _rdu = str(da.rdu)
                    _ras = str(da.ras)
                    _dataset = da.get_dataset
                    if _dataset!='-':
                        datasets.add(_dataset)
                    df.loc[len(df.index)] = [_signal_field, _iec_name, _clue_attr_name_field, _clue_attr_status_field,
                                              _func_group, _cus, _rdu, _ras, _dataset]





        datasets = list(datasets)
        datasets.sort()

        dataframe_list = list()
        for dataset in datasets: # делим датафрейм на части по датасетам
            dataframe_list.append(df[df['dataset'] == dataset])

        p1 = document.add_paragraph('Основные параметры функций')
        p1.style = 'ДОК Таблица Название'
        t1 = add_table_reports(document)

        for dataframe in dataframe_list:
            add_spec_row_table_reports(t1, ('Имя набора данных:', dataframe.iloc[0]['dataset']))

            dataframe = dataframe.sort_values('alm')
            print('вторая часть марлезонского балета')
            for row in dataframe.itertuples():
                print(row)
                row_no_index = (row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8])
                dataframe = dataframe.reset_index(drop=True)
                add_row_table_reports(t1, row_no_index)





# СОХРАНЕНИЕ ДОКУМЕНТА
    bio = io.BytesIO()
    document.save(bio) # save to memory stream
    length = bio.tell()
    print(length, '++++++++++++++++++')
    bio.seek(0) # rewind the stream

    response = HttpResponse(
            bio.getvalue(),  # use the stream's contents
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    response["Content-Disposition"] = 'attachment; filename = {0}'.format(cab+".docx")
    response["Content-Encoding"] = "UTF-8"
    response['Content-Length'] = length
    
    return response
    
