from docx import Document
from docx.shared import Inches
import io
from django.http import HttpResponse, FileResponse

import pandas as pd

from .tables import add_row_table_reports, add_table_reports, add_spec_row_table_reports, add_table_sg_sw, add_row_table_sg_sw, merge_table_sg_sw, add_row_table_sg_sw_empty, merge_table_sg_sw_header, add_row_table_sg_sw_final
from .models import Cabinets, PhDLDconnections, LogicDevices, DataObjects, LDLNconnections, LogicNodeInstantiated, \
    LNtypeObjConnections, LNobject


def return_abbr(func_group): # возвращаем абревиатуру для отчета из номера функц. группы
    if func_group == 1:
        return 'АС'
    if func_group > 1 and func_group < 13:
        return 'ПС1'
    if func_group > 12 and func_group < 20:
        return 'ПС2'
    if func_group > 19 and func_group < 23:
        return 'ОС'
    return '-'

def render_report(document, table_name, ied_cabinet, cab):

    SETTING_CDCS = ('SPG', 'ING', 'ASG', 'ENG') # ОКД для уставок
    OBJ_SETT = ('ASG', 'ING') # здесь объекты для функций
    OBJ_SW = ('SPG', 'ENG') # здесь объекты для программных переключателй
    PHASES = ('A', 'B', 'C') # для РАС
    MEASURES = ('MSQI', 'MMXU', 'MMXN') # классы измерений для отчета по mms


    df = pd.DataFrame(
        columns=['_ru_ld_name', '_ru_ln_name', '_ru_signal', '_en_ld_names', '_prefix', '_ln', '_instance',
                 '_en_signal', '_clue_attr', '_status', '_func_group', '_cus', '_rdu', '_ras', '_dataset',
                 '_sgras_name', '_dxf_signal_type', '_dxf_signal_number', 'cdc', '_ln_full_name'])
    datasets = set()

    # датафрейм для генерации уставок
    df_sg = pd.DataFrame(
        columns=['_ru_ld_name', '_ru_ln_name', '_sg_name', '_sg_desc', '_sg_spg_conds', '_cdc', '_sg_modes', 'weight'])

    # датафрейм для рас
    df_ras = pd.DataFrame(
        columns=['_ru_ld_name', '_ru_ln_name', '_sg_name', '_sg_desc', '_sg_spg_conds', '_cdc', 'weight'])


    # ищем состав лог. устройств в нем
    connections = PhDLDconnections.objects.all().filter(ied=ied_cabinet)
    for item in connections:
        _en_ld_name = str(item.ld).split('_')[0]  # отрезаем часть после подчеркивания
        ld = LogicDevices.objects.get(name=item.ld)
        _ru_ld_name = ld.fb_name.split('_')[0]  # отрезаем часть после подчеркивания
        print('*******', ld.fb_name)
        ldln_conns = LDLNconnections.objects.all().filter(ld=item.ld)
        for ldln_conn in ldln_conns:
            _ru_ln_name = str(ldln_conn.ln).split('_')[0]  # отрезаем часть после подчеркивания
            got_ln = LogicNodeInstantiated.objects.get(
                short_name=ldln_conn.ln)  # ищем тип ЛУ, чтобы вывести его объекты
            _prefix = got_ln.ln_prefix
            _ln_full_name = got_ln.full_name # только для ИЗМЕРЕНИЙ mms!
            _ln = got_ln.class_name
            _instance = got_ln.get_instance_report
            lnobj_conns = LNtypeObjConnections.objects.all().filter(ln_type=got_ln.ln_type)
            # прогоняем в цикле объекты типа логического узла
            for obj in lnobj_conns:
                _en_signal = str(obj.ln_obj).split('_')[0]  # отрезаем часть после подчеркивания
                obj_obj = LNobject.objects.get(pk=obj.ln_obj_id)
                _status = obj_obj.status
                _cdc = obj_obj.cdc
                _clue_attr = obj_obj.clue_attr
                _func_group = obj_obj.func_group
                _cus = obj_obj.cus
                _rdu = obj_obj.rdu
                _ras = obj_obj.ras
                _dataset = obj_obj.get_dataset
                _ru_signal = obj_obj.signal
                _sgras_name = obj_obj.sgras_name
                _dxf_signal_type = obj_obj.signal_type
                _dxf_signal_number = obj_obj.signal_number
                _sg_modes = obj_obj.sg_modes

                if _dataset != "-":  # если датасет не пустой добавляем строчку в датафрейм
                    datasets.add(_dataset)
                    df.loc[len(df.index)] = [_ru_ld_name, _ru_ln_name, _ru_signal, _en_ld_name, _prefix, _ln,
                                             _instance, _en_signal, _clue_attr, _status, _func_group, _cus, _rdu,
                                             _ras, _dataset, _sgras_name, _dxf_signal_type, _dxf_signal_number,
                                             _cdc, _ln_full_name]
                # print('+++++++++++++',_ru_ld_name, '/',_ru_ln_name, ':', _ru_signal )

                # датафрейм для уставок

                if str(_cdc) in SETTING_CDCS:
                    weight = 1
                    if str(_cdc) == 'SPG' or str(_cdc) == 'ENG':  # чтобы переключатели вниз отсортировались
                        weight = 2
                    df_sg.loc[len(df_sg.index)] = [_ru_ld_name, _ru_ln_name, _sgras_name, _ru_signal, _status, _cdc,
                                                   _sg_modes, weight]
                    #print('>', _ru_ld_name, _ru_ln_name, _sgras_name, _ru_signal, _status, _cdc, weight)

                # датафрейм для РАС
                if str(_ras) in ('+', 'П'):
                    rows = 1
                    #if str(_cdc) in ('ACT', 'ACD') and len(_clue_attr.split(',')) > 1:
                        #rows = 3

    # выводим таблицу с ММС
    datasets = list(datasets)
    datasets.sort()

    dataframe_list = list()
    for dataset in datasets:  # делим датафрейм на части по датасетам
        dataframe_list.append(df[df['_dataset'] == dataset])

    p1 = document.add_paragraph('Наборы данных исходящих отчетов '+table_name+' '+cab)
    p1.style = 'ДОК Таблица Название'
    t1 = add_table_reports(document)

    for dataframe in dataframe_list:
        add_spec_row_table_reports(t1, ('Имя набора данных:', dataframe.iloc[0]['_dataset']))

        # dataframe = dataframe.sort_values(by=['_func_group']) # сортируем по функциональной группе
        dataframe = dataframe.sort_values(
            by=['_func_group', '_en_ld_names','_prefix' , '_ln', '_instance', '_en_signal'])  # сортируем по функциональной группе
        # print('вторая часть марлезонского балета')
        # dataframe = dataframe.reset_index(drop=True)
        for row in dataframe.itertuples():
            print('-----------row-',row)
            ln_meas = str(row[2])
            attr = str(row[3])
            if row[6] in MEASURES: # строка измерений в таблице формируется особым способом
                if row[5]=='FLT':
                    ln_meas = 'АварРежим'
                else:
                    ln_meas = 'НормРежим'
                attr = row[20]

            row_no_index = (
            str(row[1]) + ' / ' + ln_meas + ': ' + attr, str(row[4]) + '/' + str(row[5]) + str(row[6])
            + str(row[7]) + '.' + str(row[8]), row[9], row[10], return_abbr(row[11]), row[12], row[13], row[14])

            add_row_table_reports(t1, row_no_index)

    # выводим таблицу с уставками
    if not df_sg.empty:
        p2 = document.add_paragraph('Основные параметры функций '+table_name+' '+cab)
        p2.style = 'ДОК Таблица Название'
        t2 = add_table_sg_sw(document)
        df_sg = df_sg.sort_values(by=['_ru_ld_name', '_ru_ln_name', 'weight', '_sg_name'])  # сортируем по весу '_ru_ld_name', '_ru_ln_name'
        # df_sg = df_sg.reset_index(drop=True)
        func_desc = ''  # обозначение функции и узла в строчке, меняется , когда изменяется либо функция либо узел либо то и другое
        isPG = False  # флаг наличия программных ключей  в уставках, что бы добавить примечания, относящиеся к программным ключам
        for row in df_sg.itertuples():
            if str(row[6]) in OBJ_SETT:
                if func_desc == row[1] + row[2]:
                    add_row_table_sg_sw(t2, (row[3], row[4], '-'))
                else:
                    func_desc = row[1] + row[2]
                    add_row_table_sg_sw_empty(t2, (row[1] + ': ' + row[2], '', ''))
                    merge_table_sg_sw_header(t2)
                    add_row_table_sg_sw(t2, (row[3], row[4], '-'))

            if str(row[6]) in OBJ_SW:
                isPG = True
                conds_tuple = str(row[7]).split('/')  # теперь в conds_tuple все состояния переключателя

                if func_desc != row[1] + row[2]:
                    func_desc = row[1] + row[2]
                    add_row_table_sg_sw_empty(t2, (row[1] + ': ' + row[2], '', ''))
                    merge_table_sg_sw_header(t2)
                count_item = 1  # считаем строчки с состояниями
                for item in conds_tuple:
                    if count_item == 1:
                        add_row_table_sg_sw(t2, (row[3], row[4], item.strip()))
                        #print('====================', row[3], row[4], item.strip())
                    else:
                        add_row_table_sg_sw_empty(t2, ('', '', item.strip()))
                        #print('=================+++','', '', item.strip())
                    count_item += 1
                if len(conds_tuple) > 1:  # если количество состояний больше одного, то объединяем ячейки
                    rows = t2.rows
                    merge_table_sg_sw(t2, len(rows), len(conds_tuple))
        if isPG:
            add_row_table_sg_sw_final(t2)  # добавляем финальную строчку с пояснением * - по умолчанию


# выводим таблицу РАС
# -----------------------------------------------------------------------------------------------------------
# -----------------------------------------------------------------------------------------------------------
# -----------------------------------------------------------------------------------------------------------

def word_report(request, cab):

    document = Document('base_types/templates/template.docx')
    p_cab = document.add_paragraph('Шкаф '+cab)
    p_cab.style = 'ДОК Текст 3уров 2-пункт'
    #cabinets = Cabinet.objects.all()
    #cabinets = cabinets.order_by('name')
    #for item in cabinets.iterator():
     #   add_row_table_reports(t,(item.name,'','','','','','',''))

    # ищем шкаф в базе
    cabinet = Cabinets.objects.get(name=cab)

#----------------------------------------------------------------------------------------------------------------------
# анализируем состав шкафа
    ied1 = cabinet.terminal1
    ied2 = cabinet.terminal2
    ied3 = cabinet.terminal3

    if not ied2: # если только один терминал, то обозначение в таблице ИЭУ
        table_header_name = 'ИЭУ'
        render_report(document, table_header_name, ied1, cab)
    elif ied1==ied2:
        table_header_name = 'ИЭУ1, ИЭУ2'
        render_report(document, table_header_name, ied1, cab)
    elif ied1!=ied2:
        table_header_name = 'ИЭУ1'
        render_report(document, table_header_name, ied1, cab)
        table_header_name = 'ИЭУ2'
        render_report(document, table_header_name, ied2, cab)
    if ied3:
        table_header_name = 'ИЭУ3'
        render_report(document, table_header_name, ied3, cab)
#-----------------------------------------------------------------------------------------------------------------------

    document.save('reports/doc/test.docx')
# СОХРАНЕНИЕ ДОКУМЕНТА
    #bio = io.BytesIO()
    #document.save(bio) # save to memory stream
    #length = bio.tell()
    #print(length, '++++++++++++++++++')
    #bio.seek(0) # rewind the stream
    #response = HttpResponse(
            #bio.getvalue(),  # use the stream's contents
            #content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        #)

    #response["Content-Disposition"] = 'attachment; filename = {0}'.format(cab.split(' ')[1]+".docx")
    #response["Content-Encoding"] = "UTF-8"
    #response['Content-Length'] = length
    
    return FileResponse(open('reports/doc/test.docx', 'rb')) # возвращаем сгенерированный файл dxf
    
